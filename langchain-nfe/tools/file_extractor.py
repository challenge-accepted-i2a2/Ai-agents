import fitz  # PyMuPDF
import openpyxl
import docx
from pathlib import Path
import json

def extract_text_from_file(file_path: str) -> str:
    """
    Extrai texto de arquivos PDF, XLSX ou DOCX e retorna em formato JSON válido.
    
    Args:
        file_path: Caminho do arquivo a ser processado
    
    Returns:
        String JSON com o resultado da extração
    """
    try:
        path = Path(file_path)
        
        # Validar se o arquivo existe
        if not path.exists():
            return json.dumps({
                "success": False,
                "error": "Arquivo não encontrado",
                "file_path": str(file_path),
                "content": ""
            }, ensure_ascii=False)

        # Validar extensão
        ext = path.suffix.lower()
        if ext not in [".pdf", ".xlsx", ".docx"]:
            return json.dumps({
                "success": False,
                "error": f"Extensão não suportada: {ext}",
                "file_path": str(file_path),
                "supported_extensions": [".pdf", ".xlsx", ".docx"],
                "content": ""
            }, ensure_ascii=False)

        # Extrair texto baseado na extensão
        if ext == ".pdf":
            texto = extract_pdf_text(path)
        elif ext == ".xlsx":
            texto = extract_xlsx_text(path)
        elif ext == ".docx":
            texto = extract_docx_text(path)
        
        # Limpar e validar o texto extraído
        texto_limpo = texto.strip() if texto else ""
        
        # Retornar resultado em JSON
        return json.dumps({
            "success": True,
            "file_path": str(file_path),
            "file_name": path.name,
            "file_type": ext[1:],  # Remove o ponto
            "content": texto_limpo,
            "content_length": len(texto_limpo),
            "has_content": len(texto_limpo) > 0
        }, ensure_ascii=False, indent=2)
    
    except Exception as e:
        # Capturar qualquer erro e retornar em JSON
        return json.dumps({
            "success": False,
            "error": str(e),
            "error_type": type(e).__name__,
            "file_path": str(file_path),
            "content": ""
        }, ensure_ascii=False, indent=2)


def extract_pdf_text(path: Path) -> str:
    """
    Extrai texto de arquivo PDF.
    
    Args:
        path: Path do arquivo PDF
    
    Returns:
        Texto extraído do PDF
    """
    try:
        doc = fitz.open(path)
        texto = ""
        for page_num, page in enumerate(doc, start=1):
            page_text = page.get_text()
            if page_text.strip():
                texto += f"\n--- Página {page_num} ---\n"
                texto += page_text
        doc.close()
        return texto.strip()
    except Exception as e:
        return f"Erro ao extrair texto do PDF: {str(e)}"


def extract_xlsx_text(path: Path) -> str:
    """
    Extrai texto de arquivo XLSX.
    
    Args:
        path: Path do arquivo XLSX
    
    Returns:
        Texto extraído do XLSX
    """
    try:
        wb = openpyxl.load_workbook(path, data_only=True)
        texto = ""
        
        for sheet_name in wb.sheetnames:
            sheet = wb[sheet_name]
            texto += f"\n--- Planilha: {sheet_name} ---\n"
            
            for row_idx, row in enumerate(sheet.iter_rows(values_only=True), start=1):
                # Filtrar células vazias
                row_data = [str(cell) if cell is not None else "" for cell in row]
                # Só adicionar linhas que tenham algum conteúdo
                if any(cell.strip() for cell in row_data):
                    texto += " | ".join(row_data) + "\n"
        
        wb.close()
        return texto.strip()
    except Exception as e:
        return f"Erro ao extrair texto do XLSX: {str(e)}"


def extract_docx_text(path: Path) -> str:
    """
    Extrai texto de arquivo DOCX.
    
    Args:
        path: Path do arquivo DOCX
    
    Returns:
        Texto extraído do DOCX
    """
    try:
        doc = docx.Document(path)
        
        # Extrair parágrafos
        paragrafos = [p.text for p in doc.paragraphs if p.text.strip()]
        
        # Extrair tabelas (se houver)
        tabelas_texto = []
        for table in doc.tables:
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                if any(cell for cell in row_data):
                    tabelas_texto.append(" | ".join(row_data))
        
        # Combinar parágrafos e tabelas
        texto = "\n".join(paragrafos)
        
        if tabelas_texto:
            texto += "\n\n--- Tabelas ---\n"
            texto += "\n".join(tabelas_texto)
        
        return texto.strip()
    except Exception as e:
        return f"Erro ao extrair texto do DOCX: {str(e)}"